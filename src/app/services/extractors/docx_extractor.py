import os
import logging
import docx
from typing import List, Dict, Any, Optional

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocxExtractor:
    """
    Extractor for DOCX documents.
    """
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            
        Returns:
            str: Extracted text from the DOCX.
        """
        try:
            logger.info(f"Extracting text from DOCX: {file_path}")
            
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            # Extract text using python-docx
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        tables_text.append(" | ".join(row_text))
            
            # Combine all text
            text_content = "\n\n".join(paragraphs)
            if tables_text:
                text_content += "\n\nTables:\n" + "\n".join(tables_text)
            
            logger.info(f"Extracted {len(text_content)} characters from DOCX")
            return text_content
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise